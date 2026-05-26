import mimetypes
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional, Tuple

from core.models import RawDocument

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def _extract_pdf(path: str) -> Tuple[str, int]:
    import pdfplumber
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    log.debug("PDF extracted: %d pages from %s", len(pages), path)
    return "\n\n".join(pages), len(pages)


def _extract_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    log.debug("DOCX extracted: %d paragraphs from %s", len(paragraphs), path)
    return "\n".join(paragraphs)


def _extract_html(path: str) -> str:
    from bs4 import BeautifulSoup
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    log.debug("HTML extracted from %s", path)
    return text


def _extract_csv(path: str) -> str:
    import csv
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(" | ".join(f"{k}: {v}" for k, v in row.items() if v))
    log.debug("CSV extracted: %d rows from %s", len(rows), path)
    return "\n".join(rows)

def _extract_plain(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


# ---------------------------------------------------------------------------
# Public entry-point
# ---------------------------------------------------------------------------


def extract(file_path: str, file_hash: str) -> Optional[RawDocument]:
    """
    Dispatch to the correct extractor and return a RawDocument.
    Returns None if the file type is unrecognised or extraction raises.
    """
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")
    mime_type, _ = mimetypes.guess_type(file_path)
    mime_type = mime_type or "application/octet-stream"
    page_count: Optional[int] = None

    log.info("Extracting: %s (type=%s mime=%s)", file_path, ext, mime_type)

    try:
        if ext == "pdf":
            raw_text, page_count = _extract_pdf(file_path)
            doc_type = "pdf"

        elif ext == "docx":
            raw_text = _extract_docx(file_path)
            doc_type = "docx"

        elif ext == "html":
            raw_text = _extract_html(file_path)
            doc_type = "html"

        elif ext == "csv":
            raw_text = _extract_csv(file_path)
            doc_type = "csv"

        elif ext in ("txt", "md"):
            raw_text = _extract_plain(file_path)
            doc_type = ext

        else:
            log.warning("No extractor registered for extension '.%s': %s", ext, file_path)
            return None

        char_count = len(raw_text)
        log.info(
            "Extraction complete: %s → doc_type=%s chars=%d pages=%s",
            path.name, doc_type, char_count, page_count,
        )

        return RawDocument(
            file_path=file_path,
            file_hash=file_hash,
            doc_type=doc_type,
            mime_type=mime_type,
            filename=path.name,
            raw_text=raw_text,
            page_count=page_count,
            extracted_at=datetime.now(timezone.utc).isoformat(),
        )

    except Exception:
        log.exception("Extraction failed for %s", file_path)
        return None